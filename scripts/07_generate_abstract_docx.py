"""Generate FINAL abstract v4 — trimmed to 248 words (≤250 limit).
Cuts applied vs v3:
  Methods (−2): 'Patient displacement was defined as recipient residence state'
                → 'Displacement was defined as residence state'
  Results (−5): 'traveled farther than those with licensed centers (416'
                → 'traveled farther (416'
Total removed: 7 words  |  255 → 248 words
All numbers unchanged from v3 audit (2026-03-03).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/dariorocha/Desktop/ACS_Abstract_FINAL_SUBMISSION.docx"

BURGUNDY = RGBColor(0x8B, 0x1A, 0x2E)
DARK     = RGBColor(0x2C, 0x24, 0x22)
GRAY     = RGBColor(0x6A, 0x6A, 0x6A)
GOLD     = RGBColor(0xC5, 0xA9, 0x6B)
GREEN    = RGBColor(0x1A, 0x6B, 0x2E)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def gold_line(p, color="C5A96B"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1');    b.set(qn('w:color'), color)
    pBdr.append(b); pPr.append(pBdr)

doc = Document()
for sec in doc.sections:
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.25); sec.right_margin = Inches(1.25)

# ── Track ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Free Paper  ·  Transplantation / Health Equity")
r.font.name='Calibri'; r.font.size=Pt(9.5); r.font.color.rgb=GOLD; r.bold=True

# ── Title ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run(
    "Transplant Deserts in Mexico: Organ Donation Without Local Benefit,\n"
    "Forced Migration, and Catastrophic Out-of-Pocket Costs (2007–2024)"
)
r.font.name='Calibri'; r.font.size=Pt(14); r.font.color.rgb=BURGUNDY; r.bold=True

# ── Authors ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run(
    "Damaris Rocha Castellanos · Karla Barrera Perales · Alfonso López Méndez · "
    "Edgar Méndez Dávila · Josué Roberto Lozano · Manuel Garza León"
)
r.font.name='Calibri'; r.font.size=Pt(9.5); r.font.color.rgb=DARK

# ── Meta ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
r = p.add_run("Source: CENATRA — México 2007–2024  |  Word count: 248  |  Reproducibility audit: 2026-03-03")
r.font.name='Calibri'; r.font.size=Pt(8.5); r.font.color.rgb=GRAY; r.italic=True
gold_line(p)

# ── ABSTRACT SECTIONS ─────────────────────────────────────────────────────────
# v4 cuts vs v3 (2026-03-03):
#   Methods: 'Patient displacement was defined as recipient residence state'
#            → 'Displacement was defined as residence state'  (−2 words)
#   Results: 'traveled farther than those with licensed centers (416'
#            → 'traveled farther (416'  (−5 words)
# All numbers UNCHANGED: N=2,896, 1,583 (54.7%), 27/84.4%, $873/205%, β=−7.67, R²=0.585

SECTIONS = [
    ("INTRODUCTION",
     "Mexico's liver transplant capacity is concentrated in three metropolitan centers; "
     "peripheral states function as transplant deserts, donating organs yet lacking local "
     "access. We quantified this structural inequity and financial burden."),

    ("METHODS",
     "We analyzed 2,896 liver transplants from CENATRA, Mexico's National Transplant "
     "Registry (2007–2024). Displacement was defined as residence state differing from "
     "transplant state. Double burden was defined as net organ export surplus "
     "and ≥70% of residents transplanted out-of-state. Licensed centers were identified from "
     "CENATRA's Establishment Registry (2018–2024). Distances were computed via OSRM. OOP "
     "costs were modeled using bus fares, lodging, and minimum wage benchmarks. OLS "
     "regression identified predictors of export rate."),

    ("RESULTS",
     "Three hub states held 58.1% of 86 licensed centers and performed 90.5% of transplants. "
     "Twenty-seven states (84.4%) met double burden criteria; 12 had zero licensed centers, "
     "donating 325 organs with 0 local transplants. Among 1,583 out-of-state travelers, "
     "median distance was 324 km (IQR 67–684); 14.4% traveled >1,000 km. Zero-center states "
     "traveled farther (416 vs. 160 km, p<0.001). Median OOP was $873 USD, 205% of Mexico's "
     "monthly minimum wage; 68.8% exceeded this under conservative assumptions. "
     "Double-burden states had higher poverty rates than hubs (44% vs. 22%, p=0.01). "
     "Donation volume (β=−7.67, p<0.001) and road distance (p=0.002) independently predicted "
     "export rate (R²=0.585)."),

    ("CONCLUSION",
     "Mexico's transplant deserts impose a triple burden: organ donation without local benefit, "
     "forced long-distance travel, and out-of-pocket costs exceeding two months' minimum wage, "
     "concentrated in the poorest states. Licensing programs in high-donation states where "
     "infrastructure exists is the most direct intervention to dismantle this inequity."),
]

for label, text in SECTIONS:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(label + ": ")
    r1.bold=True; r1.font.name='Calibri'; r1.font.size=Pt(10.5); r1.font.color.rgb=BURGUNDY
    r2 = p.add_run(text)
    r2.font.name='Calibri'; r2.font.size=Pt(10.5); r2.font.color.rgb=DARK

# ── Gold divider ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(10)
gold_line(p)

# ── Key message ───────────────────────────────────────────────────────────────
t = doc.add_table(rows=1, cols=1); t.style='Table Grid'
cell = t.rows[0].cells[0]; set_bg(cell, "F9F2EC")
p = cell.paragraphs[0]
p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
p.paragraph_format.left_indent=Pt(6)
r1 = p.add_run("Key message:  ")
r1.bold=True; r1.font.name='Calibri'; r1.font.size=Pt(10); r1.font.color.rgb=BURGUNDY
r2 = p.add_run(
    "Twelve Mexican states donate organs but perform zero local transplants. "
    "Residents face forced travel (median 324 km), catastrophic OOP costs "
    "($873 USD, 205% of monthly minimum wage), and structural exclusion from "
    "the transplant system they supply."
)
r2.font.name='Calibri'; r2.font.size=Pt(10); r2.font.color.rgb=DARK

# ── Confirmed values strip ─────────────────────────────────────────────────────
doc.add_paragraph()
t2 = doc.add_table(rows=2, cols=6); t2.style='Table Grid'

headers = ["N", "Displaced", "DB States", "Median dist.", "OOP (base)", "β donations"]
values  = ["2,896\n✓ confirmed", "1,583\n(54.7%)", "27/32\n(84.4%)", "324 km\nIQR 67–684", "$873 USD\n205% wage", "−7.67\np<0.001 · R²=0.585"]
notes   = ["exact match", "audit 2026-03-03", "audit 2026-03-03", "confirmed\n(orig. analysis)", "formula\nverified ✓", "confirmed\nfrom figure"]

for i, (h, v, n) in enumerate(zip(headers, values, notes)):
    cell = t2.rows[0].cells[i]; set_bg(cell, "8B1A2E")
    p = cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(h+"\n"); r1.bold=True; r1.font.name='Calibri'; r1.font.size=Pt(7.5); r1.font.color.rgb=GOLD
    r2 = p.add_run(v); r2.font.name='Calibri'; r2.font.size=Pt(8.5); r2.font.color.rgb=WHITE; r2.bold=True
    cell2 = t2.rows[1].cells[i]; set_bg(cell2, "F5F0EB")
    p2 = cell2.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3 = p2.add_run(n); r3.font.name='Calibri'; r3.font.size=Pt(7); r3.font.color.rgb=GRAY; r3.italic=True

p = doc.add_paragraph()
p.paragraph_format.space_after=Pt(2)
r = p.add_run("All values confirmed by reproducibility audit (2026-03-03)  ·  CENATRA 2007–2024  ·  ACS Clinical Congress")
r.font.name='Calibri'; r.font.size=Pt(8); r.font.color.rgb=GRAY; r.italic=True
p.alignment=WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print(f"Saved: {OUT}")

# ── Word count ────────────────────────────────────────────────────────────────
full = " ".join(label+": "+text for label, text in SECTIONS)
words = full.split()
print(f"Word count (whitespace split, incl. labels): {len(words)}")
for label, text in SECTIONS:
    print(f"  {label}: {len(text.split())} words")
print()
print("Changes from v3:")
print("  Methods: 'Patient displacement was defined as recipient residence state'")
print("           → 'Displacement was defined as residence state'  (−2 words)")
print("  Results: 'traveled farther than those with licensed centers (416'")
print("           → 'traveled farther (416'  (−5 words)")
print("  Total removed: 7 words  |  255 → 248")
print("  All numbers: UNCHANGED")
